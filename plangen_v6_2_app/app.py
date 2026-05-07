import re
from pathlib import Path
import streamlit as st
from docx import Document

from planner.core.markdown_mapper import parse_markdown_to_template_data
from planner.core.renderer import generate_docx_from_template, enrich_template_context
from planner.core.llm_client import call_llm
from planner.core.knowledge_ingestor import save_uploaded_knowledge, list_knowledge_documents
from planner.core.vector_store import build_vector_index, vector_index_status

APP_DIR = Path(__file__).parent
OUTPUT_DIR = APP_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

TEMPLATE_PATH = APP_DIR / "templates" / "AI_Center_QC_Plan_Template_20260310.docx"

st.set_page_config(page_title="PlanGen 稽查计划生成", layout="wide")
st.title("PlanGen｜V8.6 企业知识学习平台")
st.caption("支持 DeepSeek / 硅基流动 / Qwen 企业AI生成、知识库学习、向量RAG检索、自动模板映射。")

with st.sidebar:
    st.header("企业知识库")

    uploaded_kb = st.file_uploader(
        "上传知识库文件",
        type=["md", "txt", "docx", "pdf"],
        key="kb_upload"
    )

    if uploaded_kb:
        try:
            saved = save_uploaded_knowledge(uploaded_kb.name, uploaded_kb.getvalue())
            st.success(f"知识文件已导入：{saved.name}")
        except Exception as e:
            st.error(f"知识导入失败：{e}")

    if st.button("重建向量知识库"):
        with st.spinner("正在构建企业向量知识库..."):
            build_vector_index(force=True)
        st.success("企业向量知识库重建完成")

    status = vector_index_status()
    st.caption(f"知识块数量：{status.get('chunk_count', 0)}")

    docs = list_knowledge_documents()
    if docs:
        st.markdown("### 已学习知识文件")
        for d in docs:
            st.markdown(f"- {d}")

with st.expander("V8.6 企业AI配置说明", expanded=False):
    st.markdown("""
推荐配置：

硅基流动（推荐）
- LLM_BASE_URL=https://api.siliconflow.cn/v1
- LLM_MODEL=deepseek-ai/DeepSeek-V3

Embedding配置（可选）：
EMBEDDING_BASE_URL=https://api.siliconflow.cn/v1
EMBEDDING_MODEL=BAAI/bge-m3

部署环境变量：
LLM_API_KEY=
LLM_BASE_URL=
LLM_MODEL=
EMBEDDING_API_KEY=
EMBEDDING_BASE_URL=
EMBEDDING_MODEL=
""")

if not TEMPLATE_PATH.exists():
    st.error("系统内置模板缺失")
    st.stop()

input_mode = st.radio(
    "选择输入方式",
    ["AI智能生成", "粘贴 Markdown", "上传 Word"],
    horizontal=True,
    key="input_mode"
)

source_text = ""
project_name_override = ""

if input_mode == "AI智能生成":
    st.subheader("AI智能生成入口")

    project_name_override = st.text_input("项目名称", key="agent_project_name")

    llm_provider = st.selectbox(
        "AI模型",
        [
            "DeepSeek V3（推荐）",
            "Qwen",
            "OpenAI兼容模型",
        ],
        index=0,
    )

    user_prompt = st.text_area(
        "AI生成要求",
        value="请根据以下方案内容生成中心质控计划所需的标准Markdown内容。",
        height=120,
        key="agent_prompt"
    )

    protocol_text = st.text_area(
        "方案内容/Protocol摘要",
        height=260,
        key="agent_protocol_text"
    )

    uploaded_agent_docx = st.file_uploader(
        "上传方案或AI结果 Word",
        type=["docx"],
        key="agent_docx_input"
    )

    if uploaded_agent_docx:
        temp_docx = OUTPUT_DIR / uploaded_agent_docx.name
        temp_docx.write_bytes(uploaded_agent_docx.getvalue())
        doc = Document(str(temp_docx))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                if any(cells):
                    parts.append("| " + " | ".join(cells) + " |")
        protocol_text = protocol_text + "\n" + "\n".join(parts)
        st.success("Word内容读取完成")

    if st.button("调用企业AI生成稽查计划", type="primary"):
        try:
            with st.spinner("AI正在结合企业知识库进行分析..."):
                source_text = call_llm(user_prompt, protocol_text=protocol_text)
            st.session_state["agent_markdown_result"] = source_text
            st.success("企业AI内容生成完成")
        except Exception as e:
            st.error(f"AI调用失败：{e}")

    if st.session_state.get("agent_markdown_result"):
        source_text = st.session_state["agent_markdown_result"]
        with st.expander("查看AI生成Markdown"):
            st.text_area("AI生成结果", source_text, height=320)

elif input_mode == "粘贴 Markdown":
    source_text = st.text_area("粘贴AI生成Markdown", height=420)

else:
    uploaded_docx = st.file_uploader("上传AI质控结果 Word", type=["docx"])

    if uploaded_docx:
        temp_docx = OUTPUT_DIR / uploaded_docx.name
        temp_docx.write_bytes(uploaded_docx.getvalue())
        doc = Document(str(temp_docx))
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                if any(cells):
                    parts.append("| " + " | ".join(cells) + " |")

        source_text = "\n".join(parts)
        st.success("Word内容读取完成")

if source_text.strip():
    st.caption(f"已读取内容长度：{len(source_text)} 字符")

    if st.button("生成并下载稽查计划", type="primary"):
        try:
            data = parse_markdown_to_template_data(source_text)
            data = enrich_template_context(data)

            if project_name_override:
                data["PROJECT_TITLE"] = project_name_override
                data.setdefault("project", {})["name"] = project_name_override

            project_name = data.get("project", {}).get("name", "稽查计划") or data.get("PROJECT_TITLE", "稽查计划")

            safe_name = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", project_name).strip()[:60]

            output_path = OUTPUT_DIR / f"{safe_name} 稽查计划.docx"

            generate_docx_from_template(TEMPLATE_PATH, data, output_path)

            st.success("稽查计划生成成功")

            st.download_button(
                "下载稽查计划 Word",
                data=output_path.read_bytes(),
                file_name=output_path.name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"生成失败：{e}")
else:
    st.info("请选择AI智能生成、粘贴Markdown或上传Word")
