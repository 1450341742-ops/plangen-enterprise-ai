from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

PLACEHOLDERS = ["请填写", "待填写", "{{"]

FIELD_ALIASES = {
    "项目名称": ["PROJECT_TITLE", "project.name"],
    "研究题目": ["PROJECT_TITLE", "project.name"],
    "申办方": ["SPONSOR_NAME", "project.sponsor"],
    "申办者": ["SPONSOR_NAME", "project.sponsor"],
    "方案编号": ["PROJECT_CODE", "project.protocol_code"],
    "版本号": ["VERSION_NO", "project.version_no"],
    "版本日期": ["VERSION_DATE", "project.version_date"],
    "质控类型": ["AUDIT_TYPE", "project.audit_type"],
    "摘要总结": ["SUMMARY_TEXT"],
    "抽取原则": ["SAMPLING_PRINCIPLE", "RISK_SAMPLING_RULE"],
    "中心稽查风险评估病历抽取原则": ["RISK_SAMPLING_RULE", "SAMPLING_PRINCIPLE"],
    "试验药物规格": ["IMP_SPEC", "IMP_DESCRIPTION"],
    "药品规格": ["IMP_SPEC", "IMP_DESCRIPTION"],
    "法规依据补充说明": ["LAW_SUPPLEMENT"],
    "撰写人": ["AUTHOR"],
    "审批人": ["APPROVER"],
    "质控公司": ["AUDIT_COMPANY"],
    "稽查公司": ["AUDIT_COMPANY"],
}

TABLE_MAPPERS = [
    (["方案", "重点关注"], "criteria"),
    (["方案描述", "重点关注"], "process"),
    (["主要目的", "主要终点"], "endpoint"),
    (["类别", "重点关注"], "safety"),
    (["数据风险因素", "详细信息"], "risk"),
    (["筛选号", "基本情况"], "subject"),
]


def _get_value(data: Dict[str, Any], key: str) -> str:
    cur: Any = data
    for part in key.split("."):
        if isinstance(cur, dict):
            cur = cur.get(part)
        else:
            return ""
    return "" if cur is None else str(cur).strip()


def _best_value(label: str, data: Dict[str, Any]) -> str:
    clean_label = label.replace("：", "").replace(":", "").replace("\n", "").strip()
    for name, keys in FIELD_ALIASES.items():
        if name in clean_label:
            for key in keys:
                val = _get_value(data, key)
                if val:
                    return val
    for key, val in data.items():
        if isinstance(val, (str, int, float)) and key in clean_label:
            return str(val)
    return ""


def _has_placeholder(text: str) -> bool:
    return any(p in text for p in PLACEHOLDERS)


def _set_run_font(run, size_pt: int = 10, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _set_para_text(p, text: str) -> None:
    if not p.runs:
        p.add_run("")
    p.runs[0].text = text
    _set_run_font(p.runs[0], 10, False)
    for r in p.runs[1:]:
        r.text = ""


def _set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    _set_para_text(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def _header_text(table) -> str:
    if not table.rows:
        return ""
    return "|".join(c.text.strip().replace("\n", "") for c in table.rows[0].cells)


def _table_text(table) -> str:
    return "\n".join("|".join(c.text.strip() for c in row.cells) for row in table.rows)


def _clone_row(table, idx: int):
    idx = min(idx, len(table.rows) - 1)
    table._tbl.append(deepcopy(table.rows[idx]._tr))
    return table.rows[-1]


def _replace_table_rows(table, rows: List[Dict[str, Any]], mapper_type: str) -> None:
    if not rows or not table.rows:
        return
    template_idx = 1 if len(table.rows) > 1 else 0
    while len(table.rows) > template_idx + 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < template_idx + len(rows) + 1:
        _clone_row(table, template_idx)
    for row_obj, item in zip(table.rows[template_idx:template_idx + len(rows)], rows):
        if mapper_type == "criteria":
            vals = [item.get("criterion", ""), item.get("ai_focus", "")]
        elif mapper_type == "process":
            vals = [item.get("requirement", ""), item.get("focus", "")]
        elif mapper_type == "endpoint":
            vals = [item.get("objective", item.get("purpose", "")), item.get("endpoint", "")]
        elif mapper_type == "safety":
            vals = [item.get("category", ""), item.get("focus", "")]
        elif mapper_type == "risk":
            vals = [item.get("risk_factor", ""), item.get("detail", "")]
        elif mapper_type == "subject":
            vals = [item.get("subject_id", ""), item.get("summary", "")]
        else:
            vals = []
        for i, v in enumerate(vals):
            if i < len(row_obj.cells):
                _set_cell_text(row_obj.cells[i], str(v))


def _rows_for_mapper(data: Dict[str, Any], mapper_type: str) -> List[Dict[str, Any]]:
    if mapper_type == "criteria":
        return list(data.get("criteria_ai_rows", []) or []) + list(data.get("exclusion_ai_rows", []) or [])
    if mapper_type == "process":
        return list(data.get("process_requirement_rows", []) or [])
    if mapper_type == "endpoint":
        return list(data.get("primary_endpoint_rows", []) or []) + list(data.get("secondary_endpoint_rows", []) or [])
    if mapper_type == "safety":
        return list(data.get("safety_focus_rows", []) or [])
    if mapper_type == "risk":
        return list(data.get("risk_analysis_rows", []) or [])
    if mapper_type == "subject":
        return list(data.get("subject_rows", []) or [])
    return []


def auto_map_template(input_template_path: str | Path, data: Dict[str, Any], output_path: str | Path) -> Path:
    doc = Document(str(input_template_path))

    # 段落占位符替换：只替换含“请填写/待填写/{{}}”的位置，尽量不动固定文本
    for p in doc.paragraphs:
        text = p.text
        if not _has_placeholder(text):
            continue
        replacement = ""
        if "{{" in text and "}}" in text:
            new_text = text
            for key, value in data.items():
                if isinstance(value, (str, int, float)):
                    new_text = new_text.replace("{{" + key + "}}", str(value))
            if new_text != text:
                _set_para_text(p, new_text)
                continue
        label = text.split("请填写")[0].split("待填写")[0]
        replacement = _best_value(label, data)
        if replacement:
            _set_para_text(p, text.replace("请填写", replacement).replace("待填写", replacement))

    for table in doc.tables:
        table_all_text = _table_text(table)
        # 结构表自动循环
        for headers, mapper_type in TABLE_MAPPERS:
            if all(h in table_all_text for h in headers):
                rows = _rows_for_mapper(data, mapper_type)
                _replace_table_rows(table, rows, mapper_type)
                break
        # 普通字段表：左侧标签，右侧请填写
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            for i, cell in enumerate(row.cells):
                if not _has_placeholder(cell.text):
                    continue
                left_text = row.cells[i - 1].text if i > 0 else row.cells[0].text
                value = _best_value(left_text, data)
                if value:
                    _set_cell_text(cell, value)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
