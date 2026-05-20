from __future__ import annotations

import re
from pathlib import Path
from docx import Document


def _clean(text: str) -> str:
    text = "" if text is None else str(text)
    text = re.sub(r"^#{1,6}\s*", "", text.strip())
    text = re.sub(r"^[\-\*+·■▪•]\s*", "", text)
    return text.strip(" \t|：:")


def _heading(text: str) -> bool:
    t = _clean(text)
    return bool(t) and (re.match(r"^[一二三四五六七八九十]+[、.]\s*", t) is not None or re.match(r"^\d+(?:\.\d+)*\s+", t) is not None or any(t.startswith(x) for x in ["项目名称", "申办方", "申办者", "质控类型", "受试者筛选", "研究目的", "法规依据"]))


def extract_summary_fallback(source_text: str) -> str:
    lines = source_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    start = None
    for i, raw in enumerate(lines):
        if re.fullmatch(r"(?:[一二三四五六七八九十]+[、.]\s*)?摘要(?:总结)?", _clean(raw)):
            start = i + 1
            break
    if start is None:
        return ""
    out = []
    for raw in lines[start:]:
        t = _clean(raw)
        if not t:
            continue
        if re.fullmatch(r"(?:完整\s*)?mark\s*down", t, flags=re.I):
            break
        if _heading(t):
            break
        if t in {"请填写", "待填写", "xxxxx", "XXXX"}:
            continue
        out.append(t)
    return "\n".join(out).strip()


def _set_paragraph(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def ensure_summary_written(output_path: str | Path, summary_text: str) -> None:
    summary_text = (summary_text or "").strip()
    if not summary_text:
        return
    output_path = Path(output_path)
    doc = Document(str(output_path))
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if "摘要总结" not in _clean(paragraph.text):
            continue
        target = None
        for p in paragraphs[idx + 1:]:
            txt = _clean(p.text)
            if _heading(txt):
                break
            if not txt and target is None:
                target = p
            elif txt:
                return
        if target is None and idx + 1 < len(paragraphs):
            target = paragraphs[idx + 1]
        if target is not None:
            _set_paragraph(target, summary_text)
            doc.save(str(output_path))
        return
