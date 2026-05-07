from __future__ import annotations

from pathlib import Path
from docx import Document


def docx_to_markdown_text(docx_path: str | Path) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")
    return "\n".join(parts)


def _ocr_pdf_page(page) -> str:
    try:
        import pytesseract
        from PIL import Image

        pix = page.get_pixmap(matrix=None, dpi=220)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        return pytesseract.image_to_string(img, lang="chi_sim+eng") or ""
    except Exception:
        return ""


def pdf_to_text(pdf_path: str | Path, enable_ocr: bool = True) -> str:
    """Extract text from PDF. Text PDFs use PyMuPDF; scanned pages fallback to OCR if pytesseract is available."""

    pdf_path = Path(pdf_path)

    try:
        import fitz
    except Exception as e:
        raise RuntimeError("缺少PDF解析依赖 PyMuPDF，请在 requirements.txt 中加入 PyMuPDF") from e

    parts: list[str] = []

    doc = fitz.open(str(pdf_path))

    for page_no, page in enumerate(doc, 1):
        text = (page.get_text("text") or "").strip()

        # 扫描件OCR兜底
        if enable_ocr and len(text) < 30:
            ocr_text = _ocr_pdf_page(page).strip()
            if ocr_text:
                text = ocr_text

        if text:
            parts.append(f"\n\n--- PDF第{page_no}页 ---\n{text}")

    result = "\n".join(parts).strip()

    if not result:
        result = "未能从PDF中提取文本。该PDF可能为扫描件，请安装 tesseract-ocr、pytesseract、Pillow。"

    return result


def image_to_text(image_path: str | Path) -> str:
    try:
        import pytesseract
        from PIL import Image
    except Exception as e:
        raise RuntimeError("缺少OCR依赖，请安装 pytesseract 和 Pillow") from e

    img = Image.open(str(image_path))

    return pytesseract.image_to_string(img, lang="chi_sim+eng") or ""


def uploaded_file_to_text(file_path: str | Path) -> str:
    file_path = Path(file_path)

    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return docx_to_markdown_text(file_path)

    if suffix == ".pdf":
        return pdf_to_text(file_path, enable_ocr=True)

    if suffix in [".png", ".jpg", ".jpeg", ".webp"]:
        return image_to_text(file_path)

    if suffix in [".txt", ".md"]:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    raise ValueError(f"暂不支持的文件格式：{suffix}")
