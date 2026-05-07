from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List

from docx import Document

BASE_DIR = Path(__file__).resolve().parents[2]
TEMPLATE_MEMORY_DIR = BASE_DIR / "template_memory"
TEMPLATE_MEMORY_DIR.mkdir(exist_ok=True)
TEMPLATE_MEMORY_PATH = TEMPLATE_MEMORY_DIR / "template_profiles.jsonl"

PLACEHOLDER_WORDS = ["请填写", "待填写", "xxxxx", "XXXX", "{{"]
TABLE_HEADER_HINTS = ["方案", "重点关注", "方案描述", "主要目的", "主要终点", "类别", "数据风险因素", "详细信息", "筛选号", "基本情况"]


def _safe_text(text: str) -> str:
    return (text or "").replace("\n", " ").strip()


def _detect_placeholder(text: str) -> bool:
    return any(x in text for x in PLACEHOLDER_WORDS)


def _infer_field(label: str) -> str:
    label = _safe_text(label)
    mapping = {
        "项目名称": "PROJECT_TITLE",
        "研究题目": "PROJECT_TITLE",
        "申办方": "SPONSOR_NAME",
        "申办者": "SPONSOR_NAME",
        "方案编号": "PROJECT_CODE",
        "版本号": "VERSION_NO",
        "版本日期": "VERSION_DATE",
        "质控类型": "AUDIT_TYPE",
        "摘要总结": "SUMMARY_TEXT",
        "抽取原则": "RISK_SAMPLING_RULE",
        "试验药物规格": "IMP_SPEC",
        "药品规格": "IMP_SPEC",
        "法规依据": "LAW_SUPPLEMENT",
        "撰写人": "AUTHOR",
        "审批人": "APPROVER",
        "质控公司": "AUDIT_COMPANY",
    }
    for k, v in mapping.items():
        if k in label:
            return v
    return "UNKNOWN"


def _infer_table_type(text: str) -> str:
    t = _safe_text(text)
    if "方案" in t and "重点关注" in t:
        return "criteria_or_focus_table"
    if "方案描述" in t and "重点关注" in t:
        return "process_requirement_table"
    if "主要目的" in t and "主要终点" in t:
        return "endpoint_table"
    if "类别" in t and "重点关注" in t:
        return "safety_focus_table"
    if "数据风险因素" in t and "详细信息" in t:
        return "risk_analysis_table"
    if "筛选号" in t and "基本情况" in t:
        return "subject_table"
    return "unknown_table"


def learn_template_profile(template_path: str | Path, template_name: str | None = None) -> Dict[str, Any]:
    template_path = Path(template_path)
    doc = Document(str(template_path))
    profile: Dict[str, Any] = {
        "template_name": template_name or template_path.name,
        "file_name": template_path.name,
        "paragraph_placeholders": [],
        "table_placeholders": [],
        "tables": [],
        "headings": [],
    }

    for idx, p in enumerate(doc.paragraphs):
        text = _safe_text(p.text)
        if not text:
            continue
        if re.match(r"^(一、|二、|三、|\d+(\.\d+)+)", text.replace(" ", "")):
            profile["headings"].append({"index": idx, "text": text})
        if _detect_placeholder(text):
            label = text.split("请填写")[0].split("待填写")[0].split("xxxxx")[0]
            profile["paragraph_placeholders"].append({
                "paragraph_index": idx,
                "text": text,
                "label": label,
                "field": _infer_field(label),
            })

    for t_idx, table in enumerate(doc.tables):
        table_text = []
        placeholders = []
        for r_idx, row in enumerate(table.rows):
            row_text = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = _safe_text(cell.text)
                row_text.append(cell_text)
                if _detect_placeholder(cell_text):
                    left = _safe_text(row.cells[c_idx - 1].text) if c_idx > 0 else cell_text
                    placeholders.append({
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "cell_index": c_idx,
                        "label": left,
                        "placeholder_text": cell_text,
                        "field": _infer_field(left),
                    })
            table_text.append(" | ".join(row_text))
        full_text = "\n".join(table_text)
        profile["tables"].append({
            "table_index": t_idx,
            "table_type": _infer_table_type(full_text),
            "row_count": len(table.rows),
            "col_count": len(table.rows[0].cells) if table.rows else 0,
            "preview": full_text[:1000],
        })
        profile["table_placeholders"].extend(placeholders)

    return profile


def save_template_profile(profile: Dict[str, Any]) -> None:
    with TEMPLATE_MEMORY_PATH.open("a", encoding="utf-8") as f:
        f.write(json.dumps(profile, ensure_ascii=False) + "\n")


def load_template_profiles(limit: int = 50) -> List[Dict[str, Any]]:
    if not TEMPLATE_MEMORY_PATH.exists():
        return []
    rows = []
    for line in TEMPLATE_MEMORY_PATH.read_text(encoding="utf-8", errors="ignore").splitlines():
        try:
            rows.append(json.loads(line))
        except Exception:
            continue
    return rows[-limit:]


def template_learning_summary(profile: Dict[str, Any]) -> str:
    return "\n".join([
        f"模板名称：{profile.get('template_name')}",
        f"识别标题数：{len(profile.get('headings', []))}",
        f"段落占位符数：{len(profile.get('paragraph_placeholders', []))}",
        f"表格数：{len(profile.get('tables', []))}",
        f"表格占位符数：{len(profile.get('table_placeholders', []))}",
        "表格类型：" + "、".join([t.get("table_type", "") for t in profile.get("tables", [])]),
    ])
