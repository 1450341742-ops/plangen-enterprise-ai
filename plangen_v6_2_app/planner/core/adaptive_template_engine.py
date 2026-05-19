from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

KEEP_MARKS = ["保持以下文字不变", "以下文字不变", "固定内容", "不需填写", "无需填写"]
PLACEHOLDER_MARKS = ["请填写", "待填写", "xxxxx", "XXXX", "{{"]
EMPTY_BULLETS = {"·", "•", "■", "▪", "-", "—", "_"}

FIELD_ALIASES = {
    "项目/中心名称": ["PROJECT_TITLE", "project.name"],
    "项目名称": ["PROJECT_TITLE", "project.name"],
    "研究题目": ["PROJECT_TITLE", "project.name"],
    "申办方简称": ["SPONSOR_SHORT_NAME", "SPONSOR_NAME", "project.sponsor"],
    "申办方名称": ["SPONSOR_NAME", "project.sponsor"],
    "申办方": ["SPONSOR_NAME", "project.sponsor"],
    "申办者": ["SPONSOR_NAME", "project.sponsor"],
    "方案编号": ["PROJECT_CODE", "project.protocol_code"],
    "质控类型": ["AUDIT_TYPE", "project.audit_type"],
    "摘要总结": ["SUMMARY_TEXT"],
    "抽取原则": ["SAMPLING_PRINCIPLE", "RISK_SAMPLING_RULE"],
    "中心稽查风险评估病历抽取原则": ["RISK_SAMPLING_RULE", "SAMPLING_PRINCIPLE"],
    "试验药物规格/剂型/剂量/给药方式/剂量调整": ["IMP_SPEC", "IMP_DESCRIPTION"],
    "试验药物规格": ["IMP_SPEC", "IMP_DESCRIPTION"],
    "药品规格": ["IMP_SPEC", "IMP_DESCRIPTION"],
    "特别关注不良事件": ["AESI_TEXT", "SAFETY_AESI"],
    "AESI": ["AESI_TEXT", "SAFETY_AESI"],
    "法规依据补充说明": ["LAW_SUPPLEMENT"],
    "法规依据": ["LAW_SUPPLEMENT"],
    "撰写人": ["AUTHOR"],
    "审批人": ["APPROVER"],
    "质控公司": ["AUDIT_COMPANY"],
    "稽查公司": ["AUDIT_COMPANY"],
}

TABLE_TYPES = {
    "criteria": ["方案", "重点关注"],
    "process": ["方案描述", "重点关注"],
    "endpoint": ["主要目的", "主要终点"],
    "risk": ["数据风险因素", "详细信息"],
    "subject": ["筛选号", "基本情况"],
    "law": ["法规/规范名称", "条款号/章节"],
}


def _get_value(data: Dict[str, Any], key: str) -> str:
    cur: Any = data
    for part in key.split("."):
        if isinstance(cur, dict):
            cur = cur.get(part)
        else:
            return ""
    return "" if cur is None else str(cur).strip()


def _best_value(label: str, data: Dict[str, Any]) -> str:
    label = (label or "").replace("\n", "").replace("：", "").replace(":", "").replace("{{", "").replace("}}", "").strip()
    for alias, keys in FIELD_ALIASES.items():
        if alias in label:
            for key in keys:
                value = _get_value(data, key)
                if value:
                    return value
    return ""


def _is_keep_text(text: str) -> bool:
    return any(x in (text or "") for x in KEEP_MARKS)


def _has_placeholder(text: str) -> bool:
    return any(x in (text or "") for x in PLACEHOLDER_MARKS)


def _clean_generated_text(text: Any) -> str:
    text = "" if text is None else str(text)
    out = []
    for line in text.replace("\r", "\n").split("\n"):
        t = line.strip()
        if not t or t in EMPTY_BULLETS:
            continue
        if re.fullmatch(r"[·•■▪\-—_\s]+", t):
            continue
        out.append(t)
    return "\n".join(out)


def _set_run_style_like(run, size: Optional[Pt] = None, bold: Optional[bool] = None) -> None:
    run.font.name = run.font.name or "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold


def _set_para_text(p, text: str, keep_style: bool = True, cover_title: bool = False) -> None:
    text = _clean_generated_text(text)
    if not p.runs:
        p.add_run("")
    if cover_title:
        p.runs[0].text = text
        _set_run_style_like(p.runs[0], Pt(18), True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        old_size = p.runs[0].font.size
        old_bold = p.runs[0].font.bold
        p.runs[0].text = text
        if keep_style:
            _set_run_style_like(p.runs[0], old_size, old_bold)
    for r in p.runs[1:]:
        r.text = ""


def _set_cell_text(cell, text: str) -> None:
    text = _clean_generated_text(text)
    if not cell.paragraphs:
        cell.text = text
        return
    _set_para_text(cell.paragraphs[0], text, keep_style=True)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def _table_text(table) -> str:
    return "\n".join("|".join(c.text.strip() for c in row.cells) for row in table.rows)


def _is_safety_keep_table(table) -> bool:
    text = _table_text(table)
    return "保持以下文字不变" in text or ("生命体征/护理记录单" in text and "严重不良事件" in text)


def _infer_table_type(table) -> str:
    text = _table_text(table)
    if _is_safety_keep_table(table) or _is_keep_text(text):
        return "keep"
    if "随机化程序-方案描述" in text or "随机化程序-重点关注" in text:
        return "process"
    if "主要目的" in text and ("主要终点" in text or "次要终点" in text):
        return "endpoint"
    if "入选/排除标准" in text:
        return "criteria"
    for name, headers in TABLE_TYPES.items():
        if all(h in text for h in headers):
            return name
    return "unknown"


def _clone_row(table, idx: int):
    table._tbl.append(deepcopy(table.rows[idx]._tr))
    return table.rows[-1]


def _rows_for_type(data: Dict[str, Any], table_type: str) -> List[Dict[str, Any]]:
    if table_type == "criteria":
        return list(data.get("criteria_ai_rows", []) or []) + list(data.get("exclusion_ai_rows", []) or [])
    if table_type == "process":
        return list(data.get("process_requirement_rows", []) or [])
    if table_type == "endpoint":
        return list(data.get("primary_endpoint_rows", []) or []) + list(data.get("secondary_endpoint_rows", []) or [])
    if table_type == "risk":
        return list(data.get("risk_analysis_rows", []) or [])
    if table_type == "subject":
        return list(data.get("subject_rows", []) or [])
    if table_type == "law":
        return list(data.get("law_supplement_rows", []) or data.get("LAW_SUPPLEMENT_ROWS", []) or [])
    return []


def _values_for_type(item: Dict[str, Any], table_type: str) -> List[str]:
    if table_type == "criteria":
        return [item.get("criterion", ""), item.get("ai_focus", "")]
    if table_type == "process":
        return [item.get("requirement", ""), item.get("focus", "")]
    if table_type == "endpoint":
        return [item.get("objective", item.get("purpose", "")), item.get("endpoint", "")]
    if table_type == "risk":
        return [item.get("risk_factor", ""), item.get("detail", "")]
    if table_type == "subject":
        return [item.get("subject_id", ""), item.get("summary", "")]
    if table_type == "law":
        return [item.get("regulation", ""), item.get("article", ""), item.get("original_text", ""), item.get("topic", ""), item.get("applicability", "")]
    return []


def _fill_dynamic_table(table, rows: List[Dict[str, Any]], table_type: str) -> None:
    if not rows or not table.rows:
        return
    template_idx = 1 if len(table.rows) > 1 else 0
    while len(table.rows) > template_idx + 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < template_idx + 1 + len(rows):
        _clone_row(table, template_idx)
    for row_obj, item in zip(table.rows[template_idx:template_idx + len(rows)], rows):
        vals = _values_for_type(item, table_type)
        for i, val in enumerate(vals):
            if i < len(row_obj.cells):
                _set_cell_text(row_obj.cells[i], str(val))


def _replace_double_brace_text(text: str, data: Dict[str, Any]) -> str:
    def repl(m):
        key = m.group(1).strip()
        return _best_value(key, data) or ""
    return re.sub(r"\{\{([^{}]+)\}\}", repl, text)


def _find_cover_title_paragraph(doc: Document):
    for i, p in enumerate(doc.paragraphs):
        if "中心质控计划" in p.text or "稽查计划" in p.text:
            if i > 0:
                return doc.paragraphs[i - 1]
    return None


def _cell_effective_text(cell) -> str:
    return _clean_generated_text(cell.text)


def _remove_empty_rows_in_tables(doc: Document) -> None:
    for table in doc.tables:
        if _is_safety_keep_table(table):
            continue
        for row in list(table.rows)[1:]:
            row_text = " | ".join(_cell_effective_text(c) for c in row.cells)
            if not row_text.strip():
                try:
                    table._tbl.remove(row._tr)
                except Exception:
                    pass


def _clear_empty_bullet_paragraphs(doc: Document) -> None:
    # 不能彻底删除段落，只清空“只有符号”的run，避免影响模板结构。
    pattern = re.compile(r"^[\s\n\r·•■▪\-—_]+$")
    for p in doc.paragraphs:
        if pattern.fullmatch(p.text or ""):
            for r in p.runs:
                r.text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if pattern.fullmatch(p.text or ""):
                        for r in p.runs:
                            r.text = ""


def _fill_field_cells_by_left_label(doc: Document, data: Dict[str, Any]) -> None:
    for table in doc.tables:
        if _is_safety_keep_table(table):
            continue
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            left = row.cells[0].text
            value = _best_value(left, data)
            if value and (_has_placeholder(row.cells[1].text) or not row.cells[1].text.strip()):
                _set_cell_text(row.cells[1], value)


def _force_fill_sampling(doc: Document, data: Dict[str, Any]) -> None:
    value = _get_value(data, "SAMPLING_PRINCIPLE") or _get_value(data, "RISK_SAMPLING_RULE")
    if not value:
        return
    for table in doc.tables:
        if _is_safety_keep_table(table):
            continue
        for row in table.rows:
            if len(row.cells) >= 2 and "抽取原则" in row.cells[0].text:
                _set_cell_text(row.cells[1], value)


def _force_keep_safety_template(doc: Document) -> None:
    # 防止历史错误生成内容污染安全性固定表：如左侧出现入选/排除标准，恢复为空，避免继续显示错映射。
    for table in doc.tables:
        if not _is_safety_keep_table(table):
            continue
        for row in table.rows[1:]:
            left = row.cells[0].text.strip()
            if left.startswith("入选标准") or left.startswith("排除标准"):
                _set_cell_text(row.cells[0], "")
                _set_cell_text(row.cells[1], "")


def _style_table_font(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_style_like(run, run.font.size or Pt(10), run.font.bold)


def _insert_table_after_paragraph(paragraph, rows: int, cols: int):
    tbl = OxmlElement("w:tbl")
    paragraph._p.addnext(tbl)
    table = paragraph._parent.add_table(rows=rows, cols=cols)
    tbl.addnext(table._tbl)
    tbl.getparent().remove(tbl)
    return table


def _find_law_heading_paragraph(doc: Document):
    for p in doc.paragraphs:
        if "法规依据补充说明" in p.text:
            return p
    return None


def _clear_law_plain_paragraphs(doc: Document) -> None:
    """If old output wrote law entries as plain text under 2.6, clear them before inserting table."""
    heading = _find_law_heading_paragraph(doc)
    if heading is None:
        return
    found = False
    for p in doc.paragraphs:
        if p is heading:
            found = True
            continue
        if not found:
            continue
        text = p.text.strip()
        if re.match(r"^(三、|3\.1|三\s)", text):
            break
        if "｜" in text or "|" in text or "药物临床试验质量管理规范" in text or "ICH E6" in text:
            _set_para_text(p, "")


def _ensure_law_table(doc: Document, data: Dict[str, Any]) -> None:
    rows = list(data.get("law_supplement_rows", []) or data.get("LAW_SUPPLEMENT_ROWS", []) or [])
    if not rows:
        return
    headers = ["法规/规范名称", "条款号/章节", "法规原文", "对应质控主题", "本项目适用说明"]
    # Fill existing law table if a template table already exists.
    for table in doc.tables:
        if _infer_table_type(table) == "law":
            _fill_dynamic_table(table, rows, "law")
            return
    heading = _find_law_heading_paragraph(doc)
    if heading is None:
        return
    _clear_law_plain_paragraphs(doc)
    table = _insert_table_after_paragraph(heading, 2, len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h)
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    _fill_dynamic_table(table, rows, "law")
    _style_table_font(table)


def adaptive_map_template(template_path: str | Path, data: Dict[str, Any], output_path: str | Path) -> Path:
    doc = Document(str(template_path))
    project_title = _get_value(data, "PROJECT_TITLE") or _get_value(data, "project.name")

    cover_p = _find_cover_title_paragraph(doc)
    if cover_p is not None and project_title:
        _set_para_text(cover_p, project_title, cover_title=True)

    for p in doc.paragraphs:
        text = p.text
        if not _has_placeholder(text) or _is_keep_text(text):
            continue
        if "{{" in text and "}}" in text:
            new_text = _replace_double_brace_text(text, data)
            if new_text != text:
                _set_para_text(p, new_text)
                continue
        label = text.split("请填写")[0].split("待填写")[0].split("xxxxx")[0].split("XXXX")[0]
        value = _best_value(label, data)
        if value:
            new_text = text.replace("请填写", value).replace("待填写", value).replace("xxxxx", value).replace("XXXX", value)
            _set_para_text(p, new_text)

    for table in doc.tables:
        table_text = _table_text(table)
        if _is_safety_keep_table(table) or _is_keep_text(table_text):
            for row in table.rows:
                for cell in row.cells:
                    if "{{" in cell.text and "}}" in cell.text:
                        new_text = _replace_double_brace_text(cell.text, data)
                        if new_text != cell.text:
                            _set_cell_text(cell, new_text)
            continue
        t_type = _infer_table_type(table)
        rows = _rows_for_type(data, t_type)
        if rows:
            _fill_dynamic_table(table, rows, t_type)
            continue
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if not _has_placeholder(cell.text):
                    continue
                if "{{" in cell.text and "}}" in cell.text:
                    new_text = _replace_double_brace_text(cell.text, data)
                    if new_text != cell.text:
                        _set_cell_text(cell, new_text)
                        continue
                label = row.cells[idx - 1].text if idx > 0 else cell.text
                value = _best_value(label, data)
                if value:
                    _set_cell_text(cell, value)

    _ensure_law_table(doc, data)
    _fill_field_cells_by_left_label(doc, data)
    _force_fill_sampling(doc, data)
    _force_keep_safety_template(doc)
    _clear_empty_bullet_paragraphs(doc)
    _remove_empty_rows_in_tables(doc)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
