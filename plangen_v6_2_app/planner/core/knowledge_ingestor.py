from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document

BASE_DIR = Path(__file__).resolve().parents[2]
KB_DIR = BASE_DIR / "knowledge_base"
KB_DIR.mkdir(exist_ok=True)


def _safe_name(name: str) -> str:
    keep = []
    for ch in name:
        if ch.isalnum() or ch in "._-中文法规稽查质控模板规则项目报告SOP ":
            keep.append(ch)
        else:
            keep.append("_")
    return "".join(keep).strip()[:80] or "knowledge"


def extract_docx_text(path: str | Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")
    return "\n".join(parts)


def save_uploaded_knowledge(filename: str, content: bytes) -> Path:
    suffix = Path(filename).suffix.lower()
    stem = _safe_name(Path(filename).stem)
    raw_path = KB_DIR / f"{stem}{suffix}"
    raw_path.write_bytes(content)

    if suffix == ".docx":
        text = extract_docx_text(raw_path)
        md_path = KB_DIR / f"{stem}.md"
        md_path.write_text(f"# {stem}\n\n{text}", encoding="utf-8")
        return md_path

    if suffix in [".md", ".txt"]:
        return raw_path

    # PDF/其他格式暂先留存，后续可接入专用解析器。
    note_path = KB_DIR / f"{stem}.txt"
    note_path.write_text(f"已上传文件：{filename}\n当前版本暂不自动解析该格式，请转为docx/md/txt后上传。", encoding="utf-8")
    return note_path


def list_knowledge_documents() -> list[str]:
    docs = []
    for path in sorted(KB_DIR.glob("*")):
        if path.is_file() and path.name != "README.md":
            docs.append(path.name)
    return docs
