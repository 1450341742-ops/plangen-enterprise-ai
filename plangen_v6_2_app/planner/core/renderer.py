from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import re
import shutil
import subprocess
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

KEEP_UNCHANGED_MARK = "保持以下文字不变"


def enrich_template_context(data: Dict[str, Any]) -> Dict[str, Any]:
    data = dict(data)
    project = dict(data.get("project", {}))
    data["PROJECT_TITLE"] = project.get("name", data.get("PROJECT_TITLE", ""))
    data["SPONSOR_NAME"] = project.get("sponsor", data.get("SPONSOR_NAME", ""))
    data["PROJECT_CODE"] = project.get("protocol_code", data.get("PROJECT_CODE", ""))
    data["VERSION_NO"] = project.get("version_no", data.get("VERSION_NO", "V1.0"))
    data["VERSION_DATE"] = project.get("version_date", data.get("VERSION_DATE", ""))
    data["AUDIT_TYPE"] = project.get("audit_type", data.get("AUDIT_TYPE", "中心质控"))
    data.setdefault("AUTHOR", "苗田")
    data.setdefault("APPROVER", "张艳")
    data.setdefault("AUDIT_COMPANY", "北京万宁睿和医药科技有限公司")
    return data


def _safe_text(value) -> str:
    if value is None:
        return ""
    text = str(value).replace("[填写]", "待填写").replace("↵", "").replace("\r", "\n")
    lines = []
    for line in text.split("\n"):
        t = line.strip()
        if not t or t in {"■", "▪", "·", "•", "-"}:
            continue
        t = re.sub(r"^[■▪·•]\s*", "", t).strip()
        if t:
            lines.append(t)
    return "\n".join(lines).strip()


def _font_run(run, size_pt: int, bold: bool):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _set_para(p: Paragraph, text, size_pt: int = 10, bold: bool = False, align=None):
    text = _safe_text(text)
    if not p.runs:
        p.add_run("")
    p.runs[0].text = text
    _font_run(p.runs[0], size_pt, bold)
    for r in p.runs[1:]:
        r.text = ""
    if align is not None:
        p.alignment = align


def _set_cell(cell, text, size_pt: int = 10, bold: bool = False):
    if not cell.paragraphs:
        cell.text = ""
    _set_para(cell.paragraphs[0], text, size_pt, bold)
    for p in cell.paragraphs[1:]:
        p.text = ""


def _paragraph_index(doc: Document, keys: List[str]) -> Optional[int]:
    for i, p in enumerate(doc.paragraphs):
        t = p.text.replace(" ", "")
        if any(k.replace(" ", "") in t for k in keys):
            return i
    return None


def _section_is_keep(doc: Document, keys: List[str]) -> bool:
    idx = _paragraph_index(doc, keys)
    return idx is not None and KEEP_UNCHANGED_MARK in doc.paragraphs[idx].text


def _header_text(table: Table) -> str:
    if not table.rows:
        return ""
    return "|".join(c.text.strip().replace("\n", "") for c in table.rows[0].cells)


def _table_text(table: Table) -> str:
    return "\n".join("|".join(c.text.strip() for c in row.cells) for row in table.rows)


def _find_table(doc: Document, headers: List[str], occurrence: int = 1) -> Optional[Table]:
    count = 0
    for table in doc.tables:
        text = _header_text(table) + "\n" + _table_text(table)
        if all(h in text for h in headers):
            count += 1
            if count == occurrence:
                return table
    return None


def _clone_row(table: Table, idx: int):
    idx = min(idx, len(table.rows) - 1)
    table._tbl.append(deepcopy(table.rows[idx]._tr))
    return table.rows[-1]


def _replace_table_rows(table: Optional[Table], rows: List[Dict], row_builder):
    if table is None or not table.rows:
        return
    template_idx = 1 if len(table.rows) > 1 else 0
    while len(table.rows) > template_idx + 1:
        table._tbl.remove(table.rows[-1]._tr)
    if not rows:
        for cell in table.rows[template_idx].cells:
            _set_cell(cell, "")
        return
    all_rows = [table.rows[template_idx]]
    for _ in rows[1:]:
        all_rows.append(_clone_row(table, template_idx))
    for row_obj, item in zip(all_rows, rows):
        values = row_builder(item)
        for i, v in enumerate(values):
            if i < len(row_obj.cells):
                _set_cell(row_obj.cells[i], v, 10, False)


def _replace_cover_title(doc: Document, title: str):
    idx = _paragraph_index(doc, ["中心质控计划"])
    if idx is None:
        return
    target_idx = max(0, idx - 1)
    _set_para(doc.paragraphs[target_idx], title, 18, True, WD_ALIGN_PARAGRAPH.CENTER)


def _fill_labeled_value(doc: Document, label: str, value: str):
    idx = _paragraph_index(doc, [label])
    if idx is None:
        return
    p = doc.paragraphs[idx]
    normalized = p.text.replace(" ", "")
    label_norm = label.replace(" ", "")
    if normalized.rstrip().endswith(label_norm) and idx + 1 < len(doc.paragraphs):
        _set_para(doc.paragraphs[idx + 1], value, 10, False)
    else:
        _set_para(p, f"{label}{value}", 10, False)


def _placeholder_map(data: Dict) -> Dict[str, str]:
    now = datetime.now()
    criteria = (data.get("criteria_ai_rows", []) or []) + (data.get("exclusion_ai_rows", []) or [])
    process = data.get("process_requirement_rows", []) or []
    primary = data.get("primary_endpoint_rows", []) or []
    secondary = data.get("secondary_endpoint_rows", []) or []
    first_criteria = criteria[0] if criteria else {}
    first_process = process[0] if process else {}
    first_primary = primary[0] if primary else {}
    return {
        "项目/中心名称": data.get("PROJECT_TITLE", ""),
        "项目名称": data.get("PROJECT_TITLE", ""),
        "申办方名称": data.get("SPONSOR_NAME", ""),
        "申办方简称": data.get("SPONSOR_NAME", ""),
        "质控类型": data.get("AUDIT_TYPE", ""),
        "摘要总结": data.get("SUMMARY_TEXT", ""),
        "中心稽查风险评估病历抽取原则": data.get("RISK_SAMPLING_RULE", ""),
        "抽取原则": data.get("SAMPLING_PRINCIPLE", ""),
        "法规依据补充说明": data.get("LAW_SUPPLEMENT", ""),
        "撰写人": data.get("AUTHOR", ""),
        "月": str(now.month),
        "日": str(now.day),
        "入选/排除标准-方案描述": first_criteria.get("criterion", ""),
        "入选/排除标准-重点关注": first_criteria.get("ai_focus", ""),
        "随机化程序-方案描述": first_process.get("requirement", ""),
        "随机化程序-重点关注": first_process.get("focus", ""),
        "主要目的": first_primary.get("objective", ""),
        "主要终点": first_primary.get("endpoint", ""),
        "次要终点": "\n".join(f"{r.get('objective','')}：{r.get('endpoint','')}" for r in secondary),
        "特别关注不良事件（AESI）": data.get("AESI", ""),
        "试验药物规格/剂型/剂量/给药方式/剂量调整": data.get("IMP_SPEC", "") or data.get("IMP_DESCRIPTION", ""),
    }


def _replace_placeholders(doc: Document, data: Dict):
    mapping = _placeholder_map(data)
    for p in doc.paragraphs:
        if "{{" in p.text:
            text = p.text
            for key, value in mapping.items():
                text = text.replace("{{" + key + "}}", _safe_text(value))
            _set_para(p, text, 10, False)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text:
                        text = p.text
                        for key, value in mapping.items():
                            text = text.replace("{{" + key + "}}", _safe_text(value))
                        _set_para(p, text, 10, False)


def _fill_basic(doc: Document, data: Dict):
    title = data.get("PROJECT_TITLE", "")
    _replace_cover_title(doc, title)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.replace("\n", "").replace(" ", "")
            if label in {"申办方：", "申办方", "申办者：", "申办者"}:
                _set_cell(row.cells[1], data.get("SPONSOR_NAME", ""))
            elif label in {"质控公司：", "质控公司", "稽查公司：", "稽查公司"}:
                _set_cell(row.cells[1], data.get("AUDIT_COMPANY", ""))
            elif label in {"版本号/版本日期：", "版本号/版本日期"}:
                v = data.get("VERSION_NO", "V1.0")
                d = data.get("VERSION_DATE", "")
                _set_cell(row.cells[1], f"{v}/{d}" if d else v)
            elif label in {"撰写人/审批人：", "撰写人/审批人"}:
                a = data.get("AUTHOR", "")
                ap = data.get("APPROVER", "")
                _set_cell(row.cells[1], f"{a}/{ap}" if ap else a)
    for label, value in [("1.1项目名称：", title), ("1.2质控类型：", data.get("AUDIT_TYPE", "")), ("1.3申办方：", data.get("SPONSOR_NAME", ""))]:
        _fill_labeled_value(doc, label, value)


def _fill_summary(doc: Document, data: Dict):
    idx = _paragraph_index(doc, ["摘要总结"])
    if idx is not None and idx + 1 < len(doc.paragraphs) and data.get("SUMMARY_TEXT"):
        _set_para(doc.paragraphs[idx + 1], data.get("SUMMARY_TEXT"), 10, False)


def _fill_imp_table(doc: Document, data: Dict):
    spec = data.get("IMP_SPEC", "")
    if not spec:
        return
    table = _find_table(doc, ["试验药物规格"])
    if table and table.rows and len(table.rows[0].cells) >= 2:
        _set_cell(table.rows[0].cells[-1], spec, 10, False)


def _insert_law_section(doc: Document, text: str):
    if not text:
        return
    idx = _paragraph_index(doc, ["2.6法规依据补充说明", "2.6 法规依据补充说明"])
    if idx is not None and idx + 1 < len(doc.paragraphs):
        _set_para(doc.paragraphs[idx + 1], text, 10, False)


def _normalize_generated_styles(doc: Document):
    for p in doc.paragraphs:
        txt = p.text.strip().replace(" ", "")
        if not txt:
            continue
        if re.match(r"^(一、|二、|三、|\d+(\.\d+)+)", txt):
            for r in p.runs:
                if r.text:
                    _font_run(r, 16, True)
        elif "中心质控计划" not in txt:
            for r in p.runs:
                if r.text:
                    _font_run(r, 10, False)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text:
                            _font_run(r, 10, False)


def generate_docx_from_template(template_path: str | Path, data: Dict[str, Any], output_path: str | Path) -> Path:
    template_path = Path(template_path)
    output_path = Path(output_path)
    if not template_path.exists():
        raise FileNotFoundError(f"模板不存在：{template_path}")

    data = enrich_template_context(data)
    doc = Document(str(template_path))

    _replace_placeholders(doc, data)
    _fill_basic(doc, data)
    _fill_summary(doc, data)

    _replace_table_rows(_find_table(doc, ["数据风险因素", "详细信息"]), data.get("risk_analysis_rows", []) or [], lambda r: [r.get("risk_factor", ""), r.get("detail", "")])
    _replace_table_rows(_find_table(doc, ["筛选号", "基本情况"]), data.get("subject_rows", []) or [], lambda r: [r.get("subject_id", ""), r.get("summary", "")])
    _replace_table_rows(_find_table(doc, ["方案", "重点关注"], occurrence=1), (data.get("criteria_ai_rows", []) or []) + (data.get("exclusion_ai_rows", []) or []), lambda r: [r.get("criterion", ""), r.get("ai_focus", "")])
    _replace_table_rows(_find_table(doc, ["方案描述", "重点关注"], occurrence=1), data.get("process_requirement_rows", []) or [], lambda r: [r.get("requirement", ""), r.get("focus", "")])
    _replace_table_rows(_find_table(doc, ["主要目的", "主要终点"]), (data.get("primary_endpoint_rows", []) or []) + (data.get("secondary_endpoint_rows", []) or []), lambda r: [r.get("objective", r.get("purpose", "")), r.get("endpoint", "")])

    if not _section_is_keep(doc, ["2.5.3.4安全性信息处理与报告", "2.5.3.4 安全性信息处理与报告"]):
        _replace_table_rows(_find_table(doc, ["类别", "重点关注"], occurrence=1), data.get("safety_focus_rows", []) or [], lambda r: [r.get("category", ""), r.get("focus", "")])

    _fill_imp_table(doc, data)
    _insert_law_section(doc, data.get("LAW_SUPPLEMENT", ""))
    _replace_table_rows(_find_table(doc, ["姓名", "邮箱", "职位/公司"]), data.get("report_send_rows", []) or [], lambda r: [r.get("name", ""), r.get("email", ""), r.get("title_company", "")])

    _normalize_generated_styles(doc)
    _replace_cover_title(doc, data.get("PROJECT_TITLE", ""))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def convert_docx_to_pdf(docx_path: str | Path, pdf_path: str | Path | None = None, output_dir: str | Path | None = None) -> Path | None:
    docx_path = Path(docx_path)
    pdf_path = Path(output_dir) / docx_path.with_suffix(".pdf").name if output_dir else docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice")
    if not soffice:
        return None
    try:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)], check=True, capture_output=True, text=True)
    except Exception:
        return None
    return pdf_path if pdf_path.exists() else None
